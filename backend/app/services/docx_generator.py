"""
Generates a consistently formatted, ATS-friendly resume DOCX.
Fixed single-column template — same output regardless of original upload format.
Words changed/added by AI vs original are highlighted yellow.
"""
import difflib
import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ACCENT = RGBColor(0x1E, 0x40, 0xAB)   # deep blue
SUBTEXT = RGBColor(0x4B, 0x55, 0x63)  # slate grey
BODY    = RGBColor(0x11, 0x18, 0x27)  # near-black


# ── XML helpers ───────────────────────────────────────────────────────────────

def _highlight(run):
    rPr = run._r.get_or_add_rPr()
    h = OxmlElement("w:highlight")
    h.set(qn("w:val"), "yellow")
    rPr.append(h)


def _color(run, rgb: RGBColor):
    run.font.color.rgb = rgb


def _bottom_border(para, hex_color: str = "1E40AB", sz: int = 8):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), hex_color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _spacing(para, before: float = 0, after: float = 0):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)


# ── Section / line detection ──────────────────────────────────────────────────

_KNOWN_SECTIONS = re.compile(
    r"^(summary|objective|profile|experience|work experience|employment history|"
    r"education|academic|skills|technical skills|core competencies|competencies|"
    r"certifications?|projects|achievements|awards|publications|languages|"
    r"interests|references|volunteer|activities|courses?)\s*:?\s*$",
    re.IGNORECASE,
)

def _is_section(line: str) -> bool:
    s = line.strip()
    if not s or len(s) > 60:
        return False
    if _KNOWN_SECTIONS.match(s):
        return True
    if s.isupper() and 2 < len(s) < 50:
        return True
    return False


def _is_bullet(line: str) -> bool:
    return bool(re.match(r"^\s*[•\-\*–▪]\s+", line))


def _strip_bullet(line: str) -> str:
    return re.sub(r"^\s*[•\-\*–▪]\s+", "", line).strip()


def _is_role_line(line: str) -> bool:
    """Detect 'Title | Company | Date' or 'Title — Company (Date)' patterns."""
    s = line.strip()
    return bool(re.search(r"[|·•–—]", s)) and len(s) < 120


# ── Diff-highlighted paragraph ────────────────────────────────────────────────

def _diff_para(doc, orig: str, new: str, size: float = 10.5,
               bold: bool = False, italic: bool = False,
               indent: float = 0, color: RGBColor = None,
               before: float = 0, after: float = 3) -> None:
    orig_w = orig.split()
    new_w  = new.split()
    matcher = difflib.SequenceMatcher(None, orig_w, new_w, autojunk=False)

    para = doc.add_paragraph()
    _spacing(para, before, after)
    if indent:
        para.paragraph_format.left_indent = Inches(indent)

    first = True
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        words = new_w[j1:j2]
        if not words:
            continue
        chunk = ("" if first else " ") + " ".join(words)
        run = para.add_run(chunk)
        run.font.size    = Pt(size)
        run.font.bold    = bold
        run.font.italic  = italic
        _color(run, color or BODY)
        if tag in ("replace", "insert"):
            _highlight(run)
        first = False


# ── Section heading ───────────────────────────────────────────────────────────

def _section_heading(doc, text: str) -> None:
    para = doc.add_paragraph()
    _spacing(para, before=10, after=3)
    _bottom_border(para, "1E40AB", sz=8)
    run = para.add_run(text.upper())
    run.bold          = True
    run.font.size     = Pt(9)
    run.font.all_caps = True
    _color(run, ACCENT)


# ── Role sub-heading: "Title | Company | Dates" ──────────────────────────────

def _role_line(doc, line: str, orig: str, size: float = 10.5) -> None:
    # Split on | — first part = title (bold), rest = company/date (normal)
    parts = [p.strip() for p in re.split(r"[|·•–—]", line, maxsplit=1)]
    para = doc.add_paragraph()
    _spacing(para, before=6, after=1)

    # title run (bold, no highlight — it's structural not AI-written content)
    title_run = para.add_run(parts[0])
    title_run.bold       = True
    title_run.font.size  = Pt(size)
    _color(title_run, BODY)

    if len(parts) > 1:
        sep_run = para.add_run("  |  ")
        sep_run.font.size = Pt(size)
        _color(sep_run, SUBTEXT)

        detail_run = para.add_run(parts[1])
        detail_run.font.size   = Pt(size - 0.5)
        detail_run.font.italic = True
        _color(detail_run, SUBTEXT)


# ── Bullet point ──────────────────────────────────────────────────────────────

def _bullet(doc, text: str, orig: str, size: float = 10.5) -> None:
    orig_w = orig.split()
    new_w  = text.split()
    matcher = difflib.SequenceMatcher(None, orig_w, new_w, autojunk=False)

    para = doc.add_paragraph(style="List Bullet")
    _spacing(para, before=1, after=2)
    para.paragraph_format.left_indent   = Inches(0.18)
    para.paragraph_format.first_line_indent = Inches(-0.13)

    first = True
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        words = new_w[j1:j2]
        if not words:
            continue
        chunk = ("" if first else " ") + " ".join(words)
        run = para.add_run(chunk)
        run.font.size = Pt(size)
        _color(run, BODY)
        if tag in ("replace", "insert"):
            _highlight(run)
        first = False


# ── Public API ────────────────────────────────────────────────────────────────

def generate_optimized_docx(
    original_text: str,
    rewritten_text: str,
    parsed_json: dict | None = None,
) -> bytes:
    """
    Build a professionally formatted, ATS-friendly resume DOCX.

    Font size and spacing automatically tighten for candidates with 10+ years
    of experience so the document stays within 2 pages.

    Args:
        original_text:  Raw text of the original uploaded resume (for word-level diff).
        rewritten_text: AI-optimised resume text (plain text, sections separated by blank lines).
        parsed_json:    Structured data from upload-time parsing
                        (name, email, phone, linkedin, location, experience_years).
    Returns:
        DOCX as raw bytes.
    """
    parsed = parsed_json or {}

    # Decide body font size — tighter for long careers to fit 2 pages
    exp_years = parsed.get("experience_years") or 0
    body_pt   = 10.0 if exp_years >= 10 else 10.5

    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(2.0)

    doc.styles["Normal"].paragraph_format.space_after  = Pt(0)
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)

    # ── HEADER ───────────────────────────────────────────────────────────────
    name = (parsed.get("name") or "Your Name").strip()

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(name_para, after=4)
    name_run = name_para.add_run(name.upper())
    name_run.bold            = True
    name_run.font.size       = Pt(20)
    name_run.font.all_caps   = True
    _color(name_run, ACCENT)

    # Contact line
    contact_parts = []
    for field in ("email", "phone", "linkedin", "location"):
        val = parsed.get(field, "")
        if val:
            contact_parts.append(val)

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(contact_para, after=6)
        c_run = contact_para.add_run("  ·  ".join(contact_parts))
        c_run.font.size = Pt(9)
        _color(c_run, SUBTEXT)

    # Full-width blue rule
    rule = doc.add_paragraph()
    _spacing(rule, before=2, after=8)
    _bottom_border(rule, "1E40AB", sz=16)

    # ── BODY ─────────────────────────────────────────────────────────────────
    orig_lines = original_text.splitlines()
    new_lines  = rewritten_text.splitlines()

    def orig(i: int) -> str:
        return orig_lines[i].strip() if i < len(orig_lines) else ""

    # Skip leading header lines (name, email, phone, etc.) before first section
    # to avoid duplicating the header block we already rendered above.
    header_values = {v.lower().strip() for v in parsed.values() if isinstance(v, str) and v.strip()}

    def _is_header_line(line: str) -> bool:
        s = line.strip().lower()
        if not s:
            return True  # blank lines before content — skip
        # Direct match against any parsed contact field value
        if s in header_values:
            return True
        # Line is composed entirely of contact field fragments (e.g. "email | phone | linkedin")
        parts = re.split(r"[|·•·,\s]+", s)
        if all(any(p in hv or hv in p for hv in header_values) for p in parts if p):
            return True
        return False

    # Find index of first real section or bullet — skip everything before it
    # that looks like contact/header info
    i = 0
    while i < len(new_lines) and not _is_section(new_lines[i]) and _is_header_line(new_lines[i]):
        i += 1
    while i < len(new_lines):
        line    = new_lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if _is_section(line):
            _section_heading(doc, stripped)
            i += 1

        elif _is_role_line(line) and not _is_bullet(line):
            _role_line(doc, stripped, orig(i), size=body_pt)
            i += 1

        elif _is_bullet(line):
            _bullet(doc, _strip_bullet(line), _strip_bullet(orig(i)), size=body_pt)
            i += 1

        else:
            _diff_para(doc, orig(i), stripped, size=body_pt, before=1, after=3)
            i += 1

    # ── LEGEND ───────────────────────────────────────────────────────────────
    legend = doc.add_paragraph()
    _spacing(legend, before=14, after=0)
    _bottom_border(legend, "CCCCCC", sz=4)

    label = legend.add_run("Note: ")
    label.bold       = True
    label.font.size  = Pt(8)
    _color(label, SUBTEXT)

    body_note = legend.add_run(
        "Highlighted words were added or rephrased by AI to improve ATS keyword match."
    )
    body_note.font.size = Pt(8)
    _color(body_note, SUBTEXT)
    _highlight(body_note)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
